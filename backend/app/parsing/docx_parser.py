import io
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.parsing.exceptions import (
    ParsingError,
    InvalidFileFormatError,
    EmptyDocumentError,
)


def parse_docx(file_bytes: bytes) -> str:
    """
    Parse text and tabular data from a DOCX byte stream with diagnostics.
    """
    try:
        doc_stream = io.BytesIO(file_bytes)
        try:
            doc = DocxDocument(doc_stream)
        except PackageNotFoundError as e:
            raise InvalidFileFormatError(f"Invalid or corrupted DOCX package: {str(e)}")
        except Exception as e:
            raise InvalidFileFormatError(f"Failed to open DOCX document: {str(e)}")

        text_content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_content.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))

        extracted_text = "\n".join(text_content).strip()

        if not extracted_text:
            raise EmptyDocumentError("The uploaded DOCX document contains no readable text.")

        return extracted_text

    except (InvalidFileFormatError, EmptyDocumentError):
        raise
    except Exception as e:
        raise ParsingError(f"Failed to parse DOCX document: {str(e)}")
