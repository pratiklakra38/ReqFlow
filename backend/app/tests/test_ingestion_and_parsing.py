import io
import unittest
import pymupdf as fitz
from docx import Document as DocxDocument

from app.parsing.exceptions import (
    InvalidFileFormatError,
    EncryptedDocumentError,
    ScannedDocumentError,
    EmptyDocumentError,
)
from app.parsing.validator import (
    sanitize_filename,
    is_valid_pdf_magic,
    is_valid_docx_structure,
    is_valid_plain_text,
    validate_file_upload,
    MIME_PDF,
    MIME_DOCX,
    MIME_TXT,
)
from app.parsing.sanitizer import sanitize_text, enforce_text_bounds
from app.parsing.pdf_parser import parse_pdf
from app.parsing.docx_parser import parse_docx


class TestValidationAndSanitization(unittest.TestCase):

    def test_sanitize_filename_traversal_and_controls(self):
        # Path traversal checks
        self.assertEqual(sanitize_filename("../../etc/passwd.pdf"), "passwd.pdf")
        self.assertEqual(sanitize_filename("..\\..\\windows\\system32\\config.docx"), "config.docx")
        self.assertEqual(sanitize_filename("/var/tmp/specs.txt"), "specs.txt")

        # Null bytes & control characters
        self.assertEqual(sanitize_filename("test\x00file\x1b.pdf"), "testfile.pdf")

        # Empty / whitespace names
        self.assertEqual(sanitize_filename(""), "unnamed_document")
        self.assertEqual(sanitize_filename("   ...   "), "unnamed_document")
        self.assertEqual(sanitize_filename(None), "unnamed_document")

        # Filename length truncation
        long_name = "a" * 300 + ".pdf"
        sanitized = sanitize_filename(long_name)
        self.assertLessEqual(len(sanitized), 255)
        self.assertTrue(sanitized.endswith(".pdf"))

    def test_pdf_magic_bytes_detection(self):
        valid_pdf_header = b"%PDF-1.7\n1 0 obj\n<<>>\nendobj\n"
        self.assertTrue(is_valid_pdf_magic(valid_pdf_header))

        invalid_header = b"\x7fELF\x02\x01\x01\x00\x00\x00\x00"
        self.assertFalse(is_valid_pdf_magic(invalid_header))

    def test_docx_structure_detection(self):
        # Create a real DOCX in memory
        doc = DocxDocument()
        doc.add_paragraph("Sample specification")
        buf = io.BytesIO()
        doc.save(buf)
        valid_docx_bytes = buf.getvalue()

        self.assertTrue(is_valid_docx_structure(valid_docx_bytes))

        # A non-DOCX zip or plain text is invalid
        self.assertFalse(is_valid_docx_structure(b"PK\x03\x04not-a-word-doc"))
        self.assertFalse(is_valid_docx_structure(b"Hello world"))

    def test_plain_text_binary_detection(self):
        # Clean text
        self.assertTrue(is_valid_plain_text("This is plain requirements text.".encode("utf-8")))

        # Binary with null bytes disguised as text
        binary_payload = b"MZ\x90\x00\x03\x00\x00\x00\x04\x00\x00\x00\xff\xff\x00\x00"
        self.assertFalse(is_valid_plain_text(binary_payload))

    def test_validate_file_upload_rejections(self):
        # Spoofed extension: binary disguised as PDF
        with self.assertRaises(InvalidFileFormatError):
            validate_file_upload(b"MZ executable header", filename="malicious.pdf")

        # Spoofed extension: plain text disguised as DOCX
        with self.assertRaises(InvalidFileFormatError):
            validate_file_upload(b"Just plain text", filename="spoofed.docx")

        # Unsupported extension
        with self.assertRaises(InvalidFileFormatError):
            validate_file_upload(b"some data", filename="script.sh")

    def test_text_sanitizer_removes_controls_and_zero_width(self):
        # Text with non-printable control characters, zero-width spaces, and messy line endings
        raw = "User\x00 Story\x08 1:\u200B Checkout\uFEFF Process\r\n\r\n\r\n\r\nGiven \u00a0 cart is not empty  \n\n\n\nWhen user clicks pay\r\n"
        cleaned = sanitize_text(raw)

        self.assertNotIn("\x00", cleaned)
        self.assertNotIn("\x08", cleaned)
        self.assertNotIn("\u200B", cleaned)
        self.assertNotIn("\uFEFF", cleaned)
        self.assertNotIn("\r", cleaned)
        self.assertNotIn("\n\n\n", cleaned)  # Collapsed excessive newlines
        self.assertIn("User Story 1: Checkout Process", cleaned)
        self.assertIn("Given   cart is not empty", cleaned)

    def test_enforce_text_bounds(self):
        # Short text is unaffected
        short_text = "Short requirements content."
        result, truncated = enforce_text_bounds(short_text, max_chars=100)
        self.assertEqual(result, short_text)
        self.assertFalse(truncated)

        # Long text is safely truncated
        long_paragraph = "The system must process transactions within 100ms. " * 50
        result, truncated = enforce_text_bounds(long_paragraph, max_chars=200)
        self.assertTrue(truncated)
        self.assertLessEqual(len(result), 250)
        self.assertIn("truncated", result)


class TestParsersDiagnostics(unittest.TestCase):

    def test_pdf_parser_valid_text(self):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "ReqFlow Automated Ingestion Test: User shall be able to upload documents.")
        pdf_bytes = doc.tobytes()
        doc.close()

        extracted = parse_pdf(pdf_bytes)
        self.assertIn("ReqFlow Automated Ingestion Test", extracted)

    def test_pdf_parser_encrypted_rejection(self):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Encrypted content")
        pdf_bytes = doc.tobytes(encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="testpassword", owner_pw="adminpassword")
        doc.close()

        with self.assertRaises(EncryptedDocumentError):
            parse_pdf(pdf_bytes)

    def test_pdf_parser_scanned_image_detection(self):
        doc = fitz.open()
        page = doc.new_page()
        # Insert a 20x20 color pixmap with 0 selectable text
        pix = fitz.Pixmap(fitz.csRGB, (0, 0, 20, 20), 0)
        pix.clear_with(200)
        page.insert_image(page.rect, pixmap=pix)
        pdf_bytes = doc.tobytes()
        doc.close()

        with self.assertRaises(ScannedDocumentError):
            parse_pdf(pdf_bytes)

    def test_pdf_parser_empty_document(self):
        doc = fitz.open()
        page = doc.new_page()  # Blank page without text or image
        pdf_bytes = doc.tobytes()
        doc.close()

        with self.assertRaises(EmptyDocumentError):
            parse_pdf(pdf_bytes)

    def test_docx_parser_valid_and_empty(self):
        # Valid DOCX with text and tables
        doc = DocxDocument()
        doc.add_paragraph("Feature: User Authentication")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Type"
        table.cell(1, 0).text = "Email"
        table.cell(1, 1).text = "String"
        buf = io.BytesIO()
        doc.save(buf)
        valid_bytes = buf.getvalue()

        extracted = parse_docx(valid_bytes)
        self.assertIn("Feature: User Authentication", extracted)
        self.assertIn("Email | String", extracted)

        # Empty DOCX
        empty_doc = DocxDocument()
        empty_buf = io.BytesIO()
        empty_doc.save(empty_buf)
        empty_bytes = empty_buf.getvalue()

        with self.assertRaises(EmptyDocumentError):
            parse_docx(empty_bytes)


if __name__ == "__main__":
    unittest.main()
